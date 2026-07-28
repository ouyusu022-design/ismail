import streamlit as st
import pandas as pd
import os
from datetime import datetime, date
import json
from shapely.geometry import Polygon, MultiPolygon
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import io
import base64

# --- 1. إدارة ملف المستخدمين ---
USERS_FILE = "users.json"

def load_users():
    if not os.path.exists(USERS_FILE):
        default_users = {
            "admin": {
                "password": "admin123", 
                "role": "Admin", 
                "name": "المدير المسؤول"
            }
        }
        with open(USERS_FILE, "w", encoding="utf-8") as f:
            json.dump(default_users, f, ensure_ascii=False, indent=4)
        return default_users
    
    with open(USERS_FILE, "r", encoding="utf-8") as f:
        return json.load(f)

def save_users(users_dict):
    with open(USERS_FILE, "w", encoding="utf-8") as f:
        json.dump(users_dict, f, ensure_ascii=False, indent=4)

# --- 2. إعدادات وإدارة البيانات الجغرافية ---
EXCEL_FILE = "mines_permits.xlsx"
LOGO_PATH = "logo.png"

PERMIT_TYPES = {
    "رخصة الاستكشاف (Permis de Recherche)": 2,
    "رخصة التنقيب (Permis de Prospection)": 4,
    "رخصة الاستغلال (Permis d'Exploitation)": 10
}

def calculate_end_date(start_date, years):
    try:
        return start_date.replace(year=start_date.year + years)
    except ValueError:
        return start_date.replace(year=start_date.year + years, day=28)

def parse_coords(coords_str):
    try:
        points = []
        lines = str(coords_str).strip().split("\n")
        for line in lines:
            if "," in line:
                x, y = map(float, line.split(","))
                points.append((x, y))
        return points
    except:
        return []

def poly_to_coords_text(geometry):
    if geometry.is_empty:
        return ""
    
    coords_list = []
    if isinstance(geometry, Polygon):
        polys = [geometry]
    elif isinstance(geometry, MultiPolygon):
        polys = geometry.geoms
    else:
        polys = []

    for p in polys:
        for x, y in p.exterior.coords:
            coords_list.append(f"{x},{y}")
            
    return "\n".join(coords_list)

def resolve_overlap_and_trim(new_points, df, current_permit_num=None):
    if len(new_points) < 3:
        return None, 0, 0, []

    new_poly = Polygon(new_points)
    trimmed_poly = new_poly
    overlap_info = []

    for idx, row in df.iterrows():
        if current_permit_num and str(row["رقم_الرخصة"]) == str(current_permit_num):
            continue
            
        coords_str = str(row.get("الإحداثيات", ""))
        existing_points = parse_coords(coords_str)
        
        if len(existing_points) >= 3:
            existing_poly = Polygon(existing_points)
            if trimmed_poly.intersects(existing_poly):
                intersection = trimmed_poly.intersection(existing_poly)
                if intersection.area > 0:
                    overlap_info.append({
                        "رقم_الرخصة": row["رقم_الرخصة"],
                        "اسم_الشركة": row["اسم_الشركة"],
                        "مساحة_التداخل": round(intersection.area, 2)
                    })
                    trimmed_poly = trimmed_poly.difference(existing_poly)

    perim = round(trimmed_poly.length, 2)
    area = round(trimmed_poly.area, 2)
    
    return trimmed_poly, perim, area, overlap_info

def load_data():
    if not os.path.exists(EXCEL_FILE):
        df = pd.DataFrame(columns=[
            "رقم_الرخصة", 
            "اسم_الشركة", 
            "نوع_الرخصة", 
            "مدة_الصلاحية", 
            "تاريخ_البداية", 
            "تاريخ_الانتهاء", 
            "حالة_الرخصة",
            "المحيط",
            "الإحداثيات",
            "ملاحظات",
            "آخر_تعديل_بواسطة"
        ])
        df.to_excel(EXCEL_FILE, index=False)
    return pd.read_excel(EXCEL_FILE)

def save_data(df):
    df.to_excel(EXCEL_FILE, index=False)

# --- 3. إعداد الصفحة والجلسة ---
st.set_page_config(page_title="المنصة الإلكترونية لإدارة المناجم", layout="wide")
# --- كود إخفاء الأيقونات والشريط العلوي والسفلي ---
st.markdown("""
    <style>
    /* إخفاء الشريط العلوي بالأيقونات كاملاً */
    header {visibility: hidden !important;}
    #MainMenu {visibility: hidden !important;}
    
    /* إخفاء زر Manage app والتولبار السفلي */
    footer {visibility: hidden !important;}
    [data-testid="stAppToolbar"] {display: none !important;}
    </style>
""", unsafe_allow_html=True)

if "logged_in" not in st.session_state:
    st.session_state["logged_in"] = False
    st.session_state["user_info"] = None

users_db = load_users()

# --- 4. واجهة تسجيل الدخول ---
if not st.session_state["logged_in"]:
    if os.path.exists(LOGO_PATH):
        col_left, col_center, col_right = st.columns([1, 1, 1])
        with col_center:
            st.image(LOGO_PATH, use_container_width=True)

    st.markdown("<h2 style='text-align: center; color: #3B82F6;'>المنصة الإلكترونية لإدارة المناجم</h2>", unsafe_allow_html=True)
    st.write("<p style='text-align: center;'>المرجو إدخال اسم المستخدم وكلمة السر للولوج إلى النظام.</p>", unsafe_allow_html=True)
    
    with st.form("login_form"):
        username_input = st.text_input("اسم المستخدم (Nom d'utilisateur):")
        password_input = st.text_input("كلمة السر (Mot de passe):", type="password")
        submit_button = st.form_submit_button("تسجيل الدخول 🔑")
        
        if submit_button:
            if username_input in users_db and users_db[username_input]["password"] == password_input:
                user_data = users_db[username_input]
                user_data["username"] = username_input
                st.session_state["logged_in"] = True
                st.session_state["user_info"] = user_data
                st.success(f"مرحباً بك {user_data['name']}!")
                st.rerun()
            else:
                st.error("❌ اسم المستخدم أو كلمة السر غير صحيحة.")

# --- 5. واجهة التطبيق الرئيسية ---
else:
    user = st.session_state["user_info"]
    
    # --- الشريط الجانبي (Sidebar) ---
    with st.sidebar:
        st.markdown(f"""
            <div style="background-color: #1E293B; padding: 10px; border-radius: 8px; margin-bottom: 10px; border: 1px solid #334155;">
                <p style="margin: 0; font-weight: bold; color: #F8FAFC; font-size: 13px;">👤 {user['name']}</p>
                <p style="margin: 2px 0 0 0; color: #94A3B8; font-size: 11px;">الصفة: <b>{user['role']}</b> | الحساب: <b>{user.get('username', '')}</b></p>
            </div>
        """, unsafe_allow_html=True)
        
        col_btn1, col_btn2 = st.columns(2)
        with col_btn1:
            if st.button("🔄 تحديث", use_container_width=True):
                st.rerun()
        with col_btn2:
            if st.button("🚪 خروج", use_container_width=True):
                st.session_state["logged_in"] = False
                st.session_state["user_info"] = None
                st.rerun()

        st.divider()

        menu = [
            "📋 عرض والبحث حسب نوع الرخصة", 
            "🔄 إعادة تجديد رخصة (Renouvellement)",
            "📜 نموذج أمر بمهمة (Ordre de Mission)",
            "🖨️ طباعة وثيقة / شهادة رخصة",
            "📂 نماذج ووثائق للتحميل",
            "➕ إضافة رخصة جديدة لشركة", 
            "✏️ تعديل / حذف رخصة"
        ]
        
        if user.get("role") == "Admin":
            menu.append("👥 إدارة المستخدمين (إضافة وتعديل الحسابات)")
            
        menu.append("🔑 تغيير كلمة السر الخاصة بي")

        choice = st.selectbox("اختر العملية:", menu)

    # --- الجزء الرئيسي للمنصة (اللوجو الرئيسي) ---
    if os.path.exists(LOGO_PATH):
        col_left, col_center, col_right = st.columns([1, 1.5, 1])
        with col_center:
            st.image(LOGO_PATH, use_container_width=True)

    st.markdown(
        "<h1 style='text-align: center; color: #2563EB; font-weight: bold; margin-top: 10px; margin-bottom: 20px;'>المنصة الإلكترونية لإدارة المناجم</h1>", 
        unsafe_allow_html=True
    )
    st.divider()

    df = load_data()

    # --- 1. عرض والبحث ---
    if choice == "📋 عرض والبحث حسب نوع الرخصة":
        st.subheader("🔍 استعراض الشركات والصلاحيات")
        
        col1, col2 = st.columns(2)
        with col1:
            selected_type = st.selectbox(
                "اختر نوع الرخصة لفلترة جميع الشركات:", 
                ["الكل"] + list(PERMIT_TYPES.keys())
            )
        with col2:
            search_company = st.text_input("🔍 أو ابحث باسم الشركة أو رقم الرخصة:")

        filtered_df = df.copy()

        if selected_type != "الكل":
            filtered_df = filtered_df[filtered_df["نوع_الرخصة"] == selected_type]

        if search_company:
            filtered_df = filtered_df[
                filtered_df["اسم_الشركة"].astype(str).str.contains(search_company, case=False, na=False) |
                filtered_df["رقم_الرخصة"].astype(str).str.contains(search_company, case=False, na=False)
            ]

        st.dataframe(filtered_df, use_container_width=True)

        if not filtered_df.empty:
            excel_buffer = io.BytesIO()
            with pd.ExcelWriter(excel_buffer, engine='openpyxl') as writer:
                filtered_df.to_excel(writer, index=False, sheet_name='Permis_Mines')
            
            st.download_button(
                label="📥 تحميل هذه البيانات كـ ملف Excel منظم (.xlsx)",
                data=excel_buffer.getvalue(),
                file_name="permis_mines.xlsx",
                mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
            )

    # --- 2. تجديد رخصة ---
    elif choice == "🔄 إعادة تجديد رخصة (Renouvellement)":
        st.subheader("🔄 تجديد مدة صلاحية الرخصة المنجمية")
        
        if df.empty:
            st.warning("لا توجد رخص مسجلة حالياً لتجديدها.")
        else:
            option_list = [f"{i} | شركة: {row['اسم_الشركة']} - رخصة رقم: {row['رقم_الرخصة']}" for i, row in df.iterrows()]
            selected_option = st.selectbox("اختر الرخصة المراد تجديد صلاحيتها:", option_list)
            
            selected_index = int(selected_option.split(" | ")[0])
            row_data = df.loc[selected_index]
            
            st.info(f"📍 **بيانات الرخصة الحالية:** شركة: **{row_data['اسم_الشركة']}** | التاريخ الحالي للانتهاء: **{row_data['تاريخ_الانتهاء']}**")
            
            with st.form("renew_form"):
                col1, col2 = st.columns(2)
                
                with col1:
                    renew_years = st.number_input("سنوات التجديد الإضافية (Années de renouvellement):", min_value=1, max_value=10, value=2)
                    try:
                        current_end = datetime.strptime(str(row_data["تاريخ_الانتهاء"]), "%Y-%m-%d").date()
                    except:
                        current_end = date.today()
                        
                    new_renewal_start = st.date_input("تاريخ سريان التجديد:", current_end)

                with col2:
                    new_end_date = calculate_end_date(new_renewal_start, renew_years)
                    st.success(f"📅 **تاريخ الانتهاء الجديد بعد التجديد:** {new_end_date.strftime('%Y-%m-%d')}")
                    
                renew_notes = st.text_area("ملاحظات القرار أو القرار الإداري للتجديد:", value=f"تم تجديد الرخصة لمدة {renew_years} سنوات إضافية بتاريخ {date.today().strftime('%Y-%m-%d')}.")
                renew_submit = st.form_submit_button("🔄 تأكيد وحفظ التجديد")
                
                if renew_submit:
                    now_str = datetime.now().strftime("%Y-%m-%d %H:%M")
                    modified_by_info = f"{user['name']} (تجديد - {now_str})"
                    status = "صالحة (Valide)" if new_end_date >= date.today() else "منتهية (Expiré)"
                    
                    df.loc[selected_index, "تاريخ_الانتهاء"] = new_end_date.strftime("%Y-%m-%d")
                    df.loc[selected_index, "حالة_الرخصة"] = status
                    df.loc[selected_index, "ملاحظات"] = str(row_data.get("ملاحظات", "")) + f" | [تجديد]: {renew_notes}"
                    df.loc[selected_index, "آخر_تعديل_بواسطة"] = modified_by_info
                    
                    save_data(df)
                    st.success(f"✅ تم تجديد الرخصة بنجاح حتى تاريخ {new_end_date.strftime('%Y-%m-%d')}!")
                    st.rerun()

    # --- 3. إنشاء Ordre de Mission ---
    elif choice == "📜 نموذج أمر بمهمة (Ordre de Mission)":
        st.subheader("📜 إنشاء وتنزيل أمر بمهمة (Ordre de Mission)")
        st.write("قم بتعبئة المعطيات التالية لتوليد وثيقة أمر بمهمة جاهزة للطباعة بصيغة Word (.docx).")

        with st.form("om_form"):
            col1, col2 = st.columns(2)
            
            with col1:
                om_number = st.text_input("رقم الأمر بالمهمة (N° OM):", value="01/2026")
                emp_name = st.text_input("الاسم الكامل للموظف (Nom & Prénom):", value=user['name'])
                emp_grade = st.text_input("الإطار / الدرجة (Grade / Fonction):", value="Ingénieur / Cadre")
                destination = st.text_input("مكان المهمة (Destination):", placeholder="مثال: إقليم أرفود / ميدلت")

            with col2:
                transport_mean = st.selectbox(
                    "وسيلة النقل (Moyen de Transport):", 
                    ["سيارة المصلحة (Véhicule de service)", "القطار / الحافلة (Transport public)", "السيارة الشخصية (Véhicule personnel)"]
                )
                start_date_om = st.date_input("تاريخ الذهاب (Date de Départ):", date.today())
                end_date_om = st.date_input("تاريخ العودة (Date de Retour):", date.today())
                object_mission = st.text_area("موضوع المهمة (Objet de la Mission):", placeholder="مثال: معاينة ميدانية لرخصة منجمية رقم ...")

            submit_om = st.form_submit_button("⚙️ توليد وثيقة Ordre de Mission")

        if submit_om:
            if not destination or not object_mission:
                st.error("⚠️ يرجى إدخال مكان وموضوع المهمة.")
            else:
                doc = Document()

                if os.path.exists(LOGO_PATH):
                    p_logo = doc.add_paragraph()
                    p_logo.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    p_logo.add_run().add_picture(LOGO_PATH, width=Inches(1.5))

                p_head = doc.add_paragraph()
                p_head.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run_head = p_head.add_run("ROYAUME DU MAROC\nMinistère de la Transition Énergétique et du Développement Durable\nDirection Régionale des Mines\n---")
                run_head.bold = True
                run_head.font.size = Pt(11)
                run_head.font.color.rgb = RGBColor(30, 58, 138)

                p_title = doc.add_paragraph()
                p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run_title = p_title.add_run(f"\nORDRE DE MISSION N° {om_number}\n")
                run_title.bold = True
                run_title.font.size = Pt(16)
                run_title.font.color.rgb = RGBColor(30, 58, 138)

                p_body = doc.add_paragraph()
                p_body.paragraph_format.line_spacing = 1.5
                
                def add_field(label, val):
                    r1 = p_body.add_run(f"{label} : ")
                    r1.bold = True
                    r1.font.size = Pt(12)
                    r2 = p_body.add_run(f"{val}\n")
                    r2.font.size = Pt(12)

                add_field("Nom et Prénom", emp_name)
                add_field("Grade / Qualité", emp_grade)
                add_field("Objet de la mission", object_mission)
                add_field("Lieu de destination", destination)
                add_field("Moyen de transport", transport_mean)
                add_field("Date de départ", start_date_om.strftime("%d/%m/%Y"))
                add_field("Date de retour", end_date_om.strftime("%d/%m/%Y"))

                p_footer = doc.add_paragraph()
                p_footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                run_foot = p_footer.add_run(f"\nFait à ...................., le {date.today().strftime('%d/%m/%Y')}\n\nLe Directeur Régional\n\n\n\n")
                run_foot.font.size = Pt(11)

                bio = io.BytesIO()
                doc.save(bio)
                bio.seek(0)

                st.success("✅ تم إنشاء الوثيقة بنجاح!")
                st.download_button(
                    label="📥 تحميل Ordre de Mission (.docx)",
                    data=bio,
                    file_name=f"Ordre_de_Mission_{emp_name.replace(' ', '_')}.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )

    # --- 4. طباعة وثيقة ---
    elif choice == "🖨️ طباعة وثيقة / شهادة رخصة":
        st.subheader("📄 إخراج واستخراج وثيقة رسمية للرخصة")
        
        if df.empty:
            st.warning("لا توجد رخص مسجلة حالياً لإخراج وثائقها.")
        else:
            option_list = [f"{i} | شركة: {row['اسم_الشركة']} - رخصة رقم: {row['رقم_الرخصة']}" for i, row in df.iterrows()]
            selected_option = st.selectbox("اختر الرخصة المراد إخراج وثيقتها الرسمية:", option_list)
            
            selected_index = int(selected_option.split(" | ")[0])
            row_data = df.loc[selected_index]
            today_str = date.today().strftime("%Y/%m/%d")

            logo_html_tag = ""
            if os.path.exists(LOGO_PATH):
                with open(LOGO_PATH, "rb") as image_file:
                    encoded_logo = base64.b64encode(image_file.read()).decode()
                    logo_html_tag = f'<img src="data:image/png;base64,{encoded_logo}" style="max-height: 80px; max-width: 150px; object-fit: contain;" />'

            document_html = f"""
<div style="border: 3px double #1E3A8A; padding: 25px; background-color: #FFFFFF; font-family: sans-serif; color: #111827; direction: rtl; text-align: right; border-radius: 8px;">
    <div style="display: flex; justify-content: space-between; align-items: center; border-bottom: 2px solid #1E3A8A; padding-bottom: 12px; margin-bottom: 20px;">
        <div style="text-align: right;">
            <h3 style="margin: 0; color: #1E3A8A; font-size: 16px;">المملكة المغربية</h3>
            <h4 style="margin: 4px 0 0 0; color: #374151; font-size: 14px;">وزارة الانتقال الطاقي والتنمية المستدامة</h4>
            <p style="margin: 2px 0 0 0; font-size: 12px; color: #4B5563;">المديرية الجهوية للمناجم</p>
        </div>
        <div style="text-align: center;">
            {logo_html_tag}
        </div>
        <div style="text-align: left; direction: ltr;">
            <p style="margin: 0; font-weight: bold; color: #1E3A8A; font-size: 15px;">ROYAUME DU MAROC</p>
            <p style="margin: 2px 0 0 0; font-size: 11px; color: #4B5563;">Ministère de la Transition Énergétique</p>
            <p style="margin: 1px 0 0 0; font-size: 11px; color: #4B5563;">Direction des Mines</p>
        </div>
    </div>

    <div style="text-align: center; margin: 20px 0;">
        <h2 style="color: #1E3A8A; text-decoration: underline; margin-bottom: 5px; font-size: 22px;">شهادة بطاقة رخصة منجمية</h2>
        <p style="font-size: 13px; color: #6B7280; margin: 0;">(Extrait Officiel du Permis Minier)</p>
    </div>

    <table style="width: 100%; border-collapse: collapse; margin-bottom: 20px; font-size: 14px;">
        <tr style="background-color: #F3F4F6;">
            <td style="padding: 10px; border: 1px solid #D1D5DB; font-weight: bold; width: 30%;">اسم الشركة / Holder:</td>
            <td style="padding: 10px; border: 1px solid #D1D5DB;">{row_data['اسم_الشركة']}</td>
        </tr>
        <tr>
            <td style="padding: 10px; border: 1px solid #D1D5DB; font-weight: bold;">رقم الرخصة / Permit N°:</td>
            <td style="padding: 10px; border: 1px solid #D1D5DB; font-weight: bold; color: #1E3A8A;">{row_data['رقم_الرخصة']}</td>
        </tr>
        <tr style="background-color: #F3F4F6;">
            <td style="padding: 10px; border: 1px solid #D1D5DB; font-weight: bold;">نوع الرخصة / Type:</td>
            <td style="padding: 10px; border: 1px solid #D1D5DB;">{row_data['نوع_الرخصة']}</td>
        </tr>
        <tr>
            <td style="padding: 10px; border: 1px solid #D1D5DB; font-weight: bold;">مدة الصلاحية / Durée:</td>
            <td style="padding: 10px; border: 1px solid #D1D5DB;">{row_data['مدة_الصلاحية']}</td>
        </tr>
        <tr style="background-color: #F3F4F6;">
            <td style="padding: 10px; border: 1px solid #D1D5DB; font-weight: bold;">تاريخ البداية والانتهاء:</td>
            <td style="padding: 10px; border: 1px solid #D1D5DB;">من <b>{row_data['تاريخ_البداية']}</b> إلى <b>{row_data['تاريخ_الانتهاء']}</b></td>
        </tr>
        <tr>
            <td style="padding: 10px; border: 1px solid #D1D5DB; font-weight: bold;">حالة الرخصة / Statut:</td>
            <td style="padding: 10px; border: 1px solid #D1D5DB; font-weight: bold;">{row_data['حالة_الرخصة']}</td>
        </tr>
        <tr style="background-color: #F3F4F6;">
            <td style="padding: 10px; border: 1px solid #D1D5DB; font-weight: bold;">المحيط والمساحة الصافية:</td>
            <td style="padding: 10px; border: 1px solid #D1D5DB;">{row_data['المحيط']}</td>
        </tr>
        <tr>
            <td style="padding: 10px; border: 1px solid #D1D5DB; font-weight: bold;">ملاحظات إضافية:</td>
            <td style="padding: 10px; border: 1px solid #D1D5DB;">{row_data.get('ملاحظات', '-')}</td>
        </tr>
    </table>

    <div style="margin-top: 15px; font-size: 13px; background-color: #F9FAFB; padding: 10px; border-right: 4px solid #1E3A8A;">
        <b>الإحداثيات الجغرافية المسجلة (Coordonnées):</b><br/>
        <pre style="margin: 5px 0 0 0; font-family: monospace; white-space: pre-wrap;">{row_data.get('الإحداثيات', 'غير محددة')}</pre>
    </div>

    <div style="margin-top: 30px; display: flex; justify-content: space-between; align-items: flex-end;">
        <div style="font-size: 12px; color: #6B7280;">
            تاريخ الاستخراج: {today_str}<br/>
            الموظف المسؤول: {user['name']}
        </div>
        <div style="text-align: center;">
            <p style="margin: 0 0 40px 0; font-weight: bold; font-size: 13px;">خاتم وتوقيع الإدارة المعنية</p>
            <p style="margin: 0; font-size: 12px; color: #9CA3AF;">______________________</p>
        </div>
    </div>
</div>
"""
            st.write("🔍 **معاينة الوثيقة الرسمية:**")
            st.components.v1.html(document_html, height=720, scrolling=True)
            st.divider()

            st.components.v1.html(
                f"""
                <script>
                function printDoc() {{
                    var printWindow = window.open('', '', 'height=850,width=1000');
                    printWindow.document.write('<html><head><title>طباعة وثيقة رخصة</title>');
                    printWindow.document.write('<style>body {{ font-family: sans-serif; padding: 20px; }}</style>');
                    printWindow.document.write('</head><body>');
                    printWindow.document.write({json.dumps(document_html)});
                    printWindow.document.write('</body></html>');
                    printWindow.document.close();
                    printWindow.focus();
                    setTimeout(function() {{ printWindow.print(); }}, 600);
                }}
                </script>
                <div style="text-align: center;">
                    <button onclick="printDoc()" style="background-color: #1E3A8A; color: white; padding: 12px 28px; border: none; border-radius: 6px; cursor: pointer; font-size: 16px; font-weight: bold; box-shadow: 0 2px 4px rgba(0,0,0,0.2);">
                        🖨️ طباعة / حفظ الوثيقة كـ PDF
                    </button>
                </div>
                """,
                height=80
            )

    # --- 5. نماذج للتحميل ---
    elif choice == "📂 نماذج ووثائق للتحميل":
        st.subheader("📂 تحميل النماذج والوثائق الإدارية (Modèles & Formulaires)")
        st.write("يمكنك تحميل النماذج الرسمية المعتمدة لاستعمالها أو تقديمها للشركات:")

        TEMPLATES_DIR = "templates"
        if not os.path.exists(TEMPLATES_DIR):
            os.makedirs(TEMPLATES_DIR)

        documents_list = [
            {"title": "نموذج أمر بمهمة", "filename": "ordre_de_mission.docx", "desc": "نموذج Word جاهز للتعبئة (Ordre de Mission)"},
            {"title": "طلب الحصول على رخصة البحث / التنقيب", "filename": "demande_permis.docx", "desc": "صيغة Word جاهزة للتعبئة"},
            {"title": "دفتر التحملات الخاص بالاستغلال المنجمي", "filename": "cahier_de_charges.pdf", "desc": "وثيقة PDF تحتوي على الشروط والالتزامات"},
            {"title": "جدول الإحداثيات الموحد (Excel)", "filename": "canevas_coordonnees.xlsx", "desc": "ملف Excel لإدخال نقاط المحيط"}
        ]

        st.divider()
        cols = st.columns(2)
        for idx, doc in enumerate(documents_list):
            file_path = os.path.join(TEMPLATES_DIR, doc["filename"])
            with cols[idx % 2]:
                with st.container(border=True):
                    st.markdown(f"### 📄 {doc['title']}")
                    st.caption(doc['desc'])
                    
                    if os.path.exists(file_path):
                        with open(file_path, "rb") as file:
                            st.download_button(
                                label=f"📥 تحميل الملف ({doc['filename'].split('.')[-1].upper()})",
                                data=file,
                                file_name=doc["filename"],
                                mime="application/octet-stream",
                                key=f"btn_{idx}"
                            )
                    else:
                        st.warning(f"⚠️ الملف غير متوفر حالياً فـ مجلد `{TEMPLATES_DIR}`")

    # --- 6. إضافة رخصة جديدة ---
    elif choice == "➕ إضافة رخصة جديدة لشركة":
        st.subheader("➕ تسجيل رخصة جديدة مع التصحيح الأوتوماتيكي للمحيط")
        
        with st.form("add_permit_form", clear_on_submit=False):
            col1, col2 = st.columns(2)
            with col1:
                permit_num = st.text_input("رقم الرخصة (Numéro de Permis):")
                company_name = st.text_input("اسم الشركة (Nom de la Société):")
                permit_type = st.selectbox("نوع الرخصة (Type de Permis):", list(PERMIT_TYPES.keys()))
                start_date = st.date_input("تاريخ بداية الرخصة (Date de Début):", date.today())
                
            with col2:
                validity_years = PERMIT_TYPES[permit_type]
                end_date = calculate_end_date(start_date, validity_years)
                st.info(f"⏳ مدة الصلاحية: **{validity_years} سنوات**")
                st.success(f"📅 تاريخ الانتهاء: **{end_date.strftime('%Y-%m-%d')}**")

            st.divider()
            st.write("📐 **إدخال الإحداثيات الجغرافية:**")
            
            coords_input = st.text_area(
                "الإحداثيات (كل نقطة فـ سطر X,Y):",
                placeholder="500000,300000\n501000,300000\n501000,301000\n500000,301000"
            )

            notes = st.text_area("ملاحظات إضافية:")
            submit = st.form_submit_button("💾 معالجة وحفظ الرخصة")
            
            if submit:
                if permit_num and company_name:
                    points = parse_coords(coords_input)
                    trimmed_poly, perim, area, overlaps = resolve_overlap_and_trim(points, df)
                    
                    if overlaps:
                        st.warning("⚠️ **تم إكتشاف تداخل مع رخصة سابقة (تطبيق مبدأ الأسبقية):**")
                        for ov in overlaps:
                            st.write(f"🔹 تم اقتطاع المساحة المتداخلة مع الرخصة رقم **{ov['رقم_الرخصة']}** ({ov['اسم_الشركة']})")

                    final_coords_text = poly_to_coords_text(trimmed_poly) if trimmed_poly else coords_input
                    status = "صالحة (Valide)" if end_date >= date.today() else "منتهية (Expiré)"
                    
                    now_str = datetime.now().strftime("%Y-%m-%d %H:%M")
                    modified_by_info = f"{user['name']} ({now_str})"

                    new_data = pd.DataFrame([{
                        "رقم_الرخصة": permit_num,
                        "اسم_الشركة": company_name,
                        "نوع_الرخصة": permit_type,
                        "مدة_الصلاحية": f"{validity_years} سنوات",
                        "تاريخ_البداية": start_date.strftime("%Y-%m-%d"),
                        "تاريخ_الانتهاء": end_date.strftime("%Y-%m-%d"),
                        "حالة_الرخصة": status,
                        "المحيط": f"{perim} m / المساحة: {area}",
                        "الإحداثيات": final_coords_text,
                        "ملاحظات": notes,
                        "آخر_تعديل_بواسطة": modified_by_info
                    }])
                    
                    df = pd.concat([df, new_data], ignore_index=True)
                    save_data(df)
                    st.success(f"✅ تم حفظ الرخصة بواسطة {user['name']}!")
                else:
                    st.error("⚠️ عافاك دخل رقم الرخصة واسم الشركة.")

    # --- 7. تعديل أو حذف ---
    elif choice == "✏️ تعديل / حذف رخصة":
        st.subheader("✏️ تعديل معطيات رخصة أو حذفها")
        
        if df.empty:
            st.warning("لا توجد رخص مسجلة حالياً.")
        else:
            option_list = [f"{i} | شركة: {row['اسم_الشركة']} - رخصة: {row['رقم_الرخصة']}" for i, row in df.iterrows()]
            selected_option = st.selectbox("اختر الرخصة للتعديل:", option_list)
            
            selected_index = int(selected_option.split(" | ")[0])
            row_data = df.loc[selected_index]
            st.divider()
            
            col1, col2 = st.columns(2)
            with col1:
                new_permit_num = st.text_input("رقم الرخصة:", value=str(row_data["رقم_الرخصة"]))
                new_company = st.text_input("اسم الشركة:", value=str(row_data["اسم_الشركة"]))
                current_type_idx = list(PERMIT_TYPES.keys()).index(row_data["نوع_الرخصة"]) if row_data["نوع_الرخصة"] in PERMIT_TYPES else 0
                new_type = st.selectbox("نوع الرخصة:", list(PERMIT_TYPES.keys()), index=current_type_idx)
                
            with col2:
                validity_years = PERMIT_TYPES[new_type]
                try:
                    curr_start = datetime.strptime(str(row_data["تاريخ_البداية"]), "%Y-%m-%d").date()
                except:
                    curr_start = date.today()
                    
                new_start_date = st.date_input("تاريخ البداية:", curr_start)
                new_end_date = calculate_end_date(new_start_date, validity_years)
                st.info(f"📅 تاريخ الانتهاء المعدل: **{new_end_date.strftime('%Y-%m-%d')}**")

            new_coords = st.text_area("الإحداثيات:", value=str(row_data.get("الإحداثيات", "")))
            new_notes = st.text_area("ملاحظات:", value=str(row_data.get("ملاحظات", "")))
            
            col_btn1, col_btn2 = st.columns(2)
            
            with col_btn1:
                if st.button("🔄 تحديث البيانات"):
                    points = parse_coords(new_coords)
                    trimmed_poly, perim, area, overlaps = resolve_overlap_and_trim(points, df, current_permit_num=new_permit_num)
                    
                    final_coords_text = poly_to_coords_text(trimmed_poly) if trimmed_poly else new_coords
                    status = "صالحة (Valide)" if new_end_date >= date.today() else "منتهية (Expiré)"
                    
                    now_str = datetime.now().strftime("%Y-%m-%d %H:%M")
                    modified_by_info = f"{user['name']} ({now_str})"

                    df.loc[selected_index] = [
                        new_permit_num, 
                        new_company, 
                        new_type, 
                        f"{validity_years} سنوات", 
                        new_start_date.strftime("%Y-%m-%d"), 
                        new_end_date.strftime("%Y-%m-%d"), 
                        status, 
                        f"{perim} m / المساحة: {area}",
                        final_coords_text,
                        new_notes,
                        modified_by_info
                    ]
                    save_data(df)
                    st.success("✅ تم تحديث بيانات الرخصة بنجاح!")
                    st.rerun()

            with col_btn2:
                if st.button("🗑️ حذف هذه الرخصة"):
                    df = df.drop(selected_index).reset_index(drop=True)
                    save_data(df)
                    st.success("✅ تم حذف الرخصة بنجاح!")
                    st.rerun()

    # --- 8. إدارة المستخدمين ---
    elif choice == "👥 إدارة المستخدمين (إضافة وتعديل الحسابات)":
        st.subheader("👥 إدارة حسابات الموظفين والصلاحيات")
        st.write("يمكنك إضافة حساب جديد، أو إعادة تعيين كلمة السر لأي موظف نسى كلمة السر الخاصة به.")

        tab1, tab2 = st.tabs(["➕ إضافة موظف جديد", "🔑 إعادة تعيين كلمة السر لموظف (Reset Password)"])

        with tab1:
            with st.form("admin_register_form"):
                col1, col2 = st.columns(2)
                with col1:
                    new_fullname = st.text_input("الاسم الكامل للموظف (Nom et Prénom):")
                    new_username = st.text_input("اسم المستخدم الجديد (Username):")
                with col2:
                    new_role = st.selectbox("دور / صلاحية الموظف:", ["Agent", "Admin"])
                    new_pass1 = st.text_input("كلمة السر (Mot de passe):", type="password")
                    new_pass2 = st.text_input("تأكيد كلمة السر:", type="password")

                reg_button = st.form_submit_button("➕ إنشاء وحفظ الحساب")

                if reg_button:
                    if not new_fullname or not new_username or not new_pass1:
                        st.error("⚠️ يرجى ملء جميع الخانات المطلوبة.")
                    elif new_username in users_db:
                        st.error("⚠️ اسم المستخدم مستعمل بالفعل، اختر اسماً آخر.")
                    elif new_pass1 != new_pass2:
                        st.error("⚠️ كلمتا السر غير متطابقتين.")
                    else:
                        users_db[new_username] = {
                            "password": new_pass1,
                            "role": new_role,
                            "name": new_fullname
                        }
                        save_users(users_db)
                        st.success(f"✅ تم إنشاء حساب الموظف ({new_fullname}) بنجاح بصلاحية {new_role}!")
                        st.rerun()

        with tab2:
            st.warning("⚠️ اختر الحساب المراد تغيير كلمة السر له وأدخل كلمة السر الجديدة مباشرة.")
            user_options = [f"{k} | ({v.get('name')})" for k, v in users_db.items()]
            selected_user_opt = st.selectbox("اختر حساب الموظف:", user_options)
            
            selected_uname = selected_user_opt.split(" | ")[0]

            with st.form("reset_pass_form"):
                admin_new_pass = st.text_input("كلمة السر الجديدة للحساب:", type="password")
                admin_confirm_pass = st.text_input("تأكيد كلمة السر الجديدة:", type="password")
                
                reset_btn = st.form_submit_button("🔄 تغيير وتحديث كلمة السر")

                if reset_btn:
                    if not admin_new_pass:
                        st.error("⚠️ يرجى إدخال كلمة السر الجديدة.")
                    elif admin_new_pass != admin_confirm_pass:
                        st.error("⚠️ كلمتا السر غير متطابقتين.")
                    else:
                        users_db[selected_uname]["password"] = admin_new_pass
                        save_users(users_db)
                        st.success(f"✅ تم تحديث كلمة السر للحساب ({selected_uname}) بنجاح!")

        st.divider()
        st.subheader("📋 قائمة الموظفين المسجلين حالياً:")
        user_list_data = [{"اسم المستخدم": k, "الاسم الكامل": v.get("name"), "الصفة/Role": v.get("role")} for k, v in users_db.items()]
        st.table(pd.DataFrame(user_list_data))

    # --- 9. تغيير كلمة السر الشخصية ---
    elif choice == "🔑 تغيير كلمة السر الخاصة بي":
        st.subheader("🔑 تغيير كلمة السر للحساب الحالي")
        
        with st.form("change_pass_form"):
            old_pass = st.text_input("كلمة السر الحالية:", type="password")
            new_pass = st.text_input("كلمة السر الجديدة:", type="password")
            confirm_pass = st.text_input("تأكيد كلمة السر الجديدة:", type="password")
            
            submit_pass = st.form_submit_button("💾 حفظ كلمة السر الجديدة")
            
            if submit_pass:
                current_username = user.get("username")
                if users_db.get(current_username, {}).get("password") != old_pass:
                    st.error("❌ كلمة السر الحالية غير صحيحة.")
                elif new_pass != confirm_pass:
                    st.error("⚠️ كلمتا السر غير متطابقتين.")
                elif not new_pass:
                    st.error("⚠️ يرجى إدخال كلمة سر جديدة.")
                else:
                    users_db[current_username]["password"] = new_pass
                    save_users(users_db)
                    st.success("✅ تم تغيير كلمة السر بنجاح!")
